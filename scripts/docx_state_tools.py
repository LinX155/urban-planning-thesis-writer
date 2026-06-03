"""Windows-friendly DOCX state, snapshot, and diff helpers.

The skill uses this script as a deterministic helper. It preserves user edits
by snapshotting DOCX text and producing lightweight diffs before any write pass.
"""

from __future__ import annotations

import argparse
import difflib
import json
import shutil
import sys
from datetime import datetime
from pathlib import Path


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def ensure_docx():
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise SystemExit("Missing dependency: python-docx. Run init_thesis_workspace.ps1 first.") from exc
    return docx


def timestamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")


def state_root(workspace: Path) -> Path:
    return workspace / ".urban-planning-thesis-writer"


def resolve_workspace_path(workspace: Path | None, raw_path: str) -> Path:
    path = Path(raw_path)
    if not path.is_absolute() and workspace is not None:
        path = workspace / path
    return path.resolve()


def extract_docx_text(docx_path: Path) -> str:
    docx = ensure_docx()
    document = docx.Document(str(docx_path))
    lines: list[str] = []
    for para in document.paragraphs:
        style = para.style.name if para.style is not None else ""
        text = para.text.strip()
        if text:
            lines.append(f"[{style}] {text}")
    for table_idx, table in enumerate(document.tables, start=1):
        lines.append(f"[TABLE {table_idx}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    for section_idx, section in enumerate(document.sections, start=1):
        header_lines = [para.text.strip() for para in section.header.paragraphs if para.text.strip()]
        footer_lines = [para.text.strip() for para in section.footer.paragraphs if para.text.strip()]
        if header_lines:
            lines.append(f"[SECTION {section_idx} HEADER]")
            lines.extend(header_lines)
        if footer_lines:
            lines.append(f"[SECTION {section_idx} FOOTER]")
            lines.extend(footer_lines)
    return "\n".join(lines) + "\n"


def snapshot(workspace: Path, docx_path: Path, label: str | None = None) -> Path:
    root = state_root(workspace)
    snap_dir = root / "state" / "snapshots"
    backup_dir = root / "state" / "backups"
    snap_dir.mkdir(parents=True, exist_ok=True)
    backup_dir.mkdir(parents=True, exist_ok=True)
    stamp = timestamp()
    safe_label = label or docx_path.stem
    text_path = snap_dir / f"{stamp}-{safe_label}.txt"
    backup_path = backup_dir / f"{stamp}-{docx_path.name}"
    text_path.write_text(extract_docx_text(docx_path), encoding="utf-8")
    shutil.copy2(docx_path, backup_path)

    progress_path = root / "state" / "progress.json"
    progress = {}
    if progress_path.exists():
        progress = load_json(progress_path)
    progress["last_snapshot"] = str(text_path)
    progress["last_backup"] = str(backup_path)
    progress_path.write_text(json.dumps(progress, ensure_ascii=False, indent=2), encoding="utf-8")
    print(text_path)
    return text_path


def diff_snapshots(workspace: Path, before: Path, after: Path, label: str | None = None) -> Path:
    root = state_root(workspace)
    diff_dir = root / "state" / "diffs"
    diff_dir.mkdir(parents=True, exist_ok=True)
    before_lines = before.read_text(encoding="utf-8").splitlines()
    after_lines = after.read_text(encoding="utf-8").splitlines()
    diff = difflib.unified_diff(
        before_lines,
        after_lines,
        fromfile=before.name,
        tofile=after.name,
        lineterm="",
        n=3,
    )
    out = diff_dir / f"{timestamp()}-{label or 'docx-diff'}.diff"
    out.write_text("\n".join(diff) + "\n", encoding="utf-8")
    print(out)
    return out


def summarize_diff(diff_path: Path, max_changes: int = 80, output_path: Path | None = None) -> str:
    added: list[str] = []
    removed: list[str] = []
    for line in diff_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        if line.startswith("+++") or line.startswith("---") or line.startswith("@@"):
            continue
        if line.startswith("+") and len(added) < max_changes:
            added.append(line[1:].strip())
        elif line.startswith("-") and len(removed) < max_changes:
            removed.append(line[1:].strip())
    summary = {
        "diff": str(diff_path),
        "removed_samples": [x for x in removed if x][:max_changes],
        "added_samples": [x for x in added if x][:max_changes],
        "guardrail": "Treat these as evidence of user edits, not as automatic global preferences. Generalize only stable terminology, structure, or academic style choices.",
    }
    text = json.dumps(summary, ensure_ascii=False, indent=2)
    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(text, encoding="utf-8")
    print(text)
    return text


def main() -> int:
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="command", required=True)

    p_snap = sub.add_parser("snapshot")
    p_snap.add_argument("--workspace", required=True)
    p_snap.add_argument("--docx", required=True)
    p_snap.add_argument("--label")

    p_diff = sub.add_parser("diff")
    p_diff.add_argument("--workspace", required=True)
    p_diff.add_argument("--before", required=True)
    p_diff.add_argument("--after", required=True)
    p_diff.add_argument("--label")

    p_sum = sub.add_parser("summarize-diff")
    p_sum.add_argument("--workspace")
    p_sum.add_argument("--diff", required=True)
    p_sum.add_argument("--output")

    args = parser.parse_args()
    if args.command == "snapshot":
        workspace = Path(args.workspace).resolve()
        snapshot(workspace, resolve_workspace_path(workspace, args.docx), args.label)
    elif args.command == "diff":
        workspace = Path(args.workspace).resolve()
        diff_snapshots(
            workspace,
            resolve_workspace_path(workspace, args.before),
            resolve_workspace_path(workspace, args.after),
            args.label,
        )
    elif args.command == "summarize-diff":
        workspace = Path(args.workspace).resolve() if args.workspace else None
        summarize_diff(
            resolve_workspace_path(workspace, args.diff),
            output_path=resolve_workspace_path(workspace, args.output) if args.output else None,
        )
    return 0


if __name__ == "__main__":
    sys.exit(main())
