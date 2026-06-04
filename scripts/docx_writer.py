"""Controlled DOCX writing helper for the thesis skill.

This script exposes a small set of reusable editing primitives for plan-guided
DOCX revision. The calling agent should choose the smallest authorized edit
surface that preserves surrounding layout and non-text anchors whenever
possible, rather than treating these commands as a fixed scenario checklist.
"""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Iterable


def ensure_docx():
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise SystemExit("Missing dependency: python-docx. Run /UPTW-plan first to bootstrap the workspace.") from exc
    return docx


def timestamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")


def state_root(workspace: Path) -> Path:
    return workspace / ".urban-planning-thesis-writer"


def resolve_workspace_path(workspace: Path, raw_path: str) -> Path:
    path = Path(raw_path)
    if not path.is_absolute():
        path = workspace / path
    return path.resolve()


def backup_docx(workspace: Path, docx_path: Path, label: str) -> Path:
    backup_dir = state_root(workspace) / "state" / "backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    out = backup_dir / f"{timestamp()}-{label}-{docx_path.name}"
    shutil.copy2(docx_path, out)
    return out


def style_names(document) -> set[str]:
    return {style.name for style in document.styles}


def pick_style(document, preferred: Iterable[str], fallback: str | None = None) -> str | None:
    names = style_names(document)
    for name in preferred:
        if name in names:
            return name
    return fallback if fallback in names else None


def heading_style(document, level: int) -> str | None:
    return pick_style(document, [f"Heading {level}", f"标题 {level}", f"标题{level}"])


def paragraph_style(document) -> str | None:
    return pick_style(document, ["Normal", "正文"])


def bullet_style(document) -> str | None:
    return pick_style(document, ["List Bullet", "项目符号", "列表项目"])


def normalized_text(text: str) -> str:
    return " ".join(text.split())


def heading_level_from_name(style_name: str) -> int | None:
    compact = "".join(style_name.split()).lower()
    if compact.startswith("heading"):
        suffix = compact.removeprefix("heading")
        if suffix.isdigit():
            return int(suffix)
    if compact.startswith("标题"):
        suffix = compact.removeprefix("标题")
        if suffix.isdigit():
            return int(suffix)
    return None


def heading_level_for_paragraph(paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    return heading_level_from_name(style_name)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def paragraph_has_page_break(paragraph) -> bool:
    return bool(
        paragraph._p.xpath(
            ".//*[local-name()='br' and @*[local-name()='type']='page'] | .//*[local-name()='lastRenderedPageBreak']"
        )
    )


def paragraph_guard_flags(paragraph) -> dict[str, bool]:
    return {
        "has_drawing": paragraph_has_drawing(paragraph),
        "has_page_break": paragraph_has_page_break(paragraph),
    }


def paragraph_is_protected(paragraph) -> bool:
    flags = paragraph_guard_flags(paragraph)
    return flags["has_drawing"] or flags["has_page_break"]


def first_text_run(paragraph):
    for run in paragraph.runs:
        if run.text or run.style or run.bold is not None or run.italic is not None or run.underline is not None:
            return run
    return None


def copy_run_properties(source_run, target_run) -> None:
    if source_run is None:
        return
    docx = ensure_docx()
    qn = docx.oxml.ns.qn
    source_rpr = source_run._r.find(qn("w:rPr"))
    target_rpr = target_run._r.find(qn("w:rPr"))
    if target_rpr is not None:
        target_run._r.remove(target_rpr)
    if source_rpr is not None:
        target_run._r.insert(0, deepcopy(source_rpr))


def replace_paragraph_properties(source_paragraph, target_paragraph) -> None:
    docx = ensure_docx()
    qn = docx.oxml.ns.qn
    source_ppr = source_paragraph._p.find(qn("w:pPr"))
    target_ppr = target_paragraph._p.find(qn("w:pPr"))
    if target_ppr is not None:
        target_paragraph._p.remove(target_ppr)
    if source_ppr is not None:
        target_paragraph._p.insert(0, deepcopy(source_ppr))


def clear_paragraph_content(paragraph) -> None:
    docx = ensure_docx()
    qn = docx.oxml.ns.qn
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def insert_paragraph_after(paragraph):
    docx = ensure_docx()
    new_p = docx.oxml.OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return docx.text.paragraph.Paragraph(new_p, paragraph._parent)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def parse_text_blocks(text: str) -> list[tuple[str, str, int | None]]:
    """Return (kind, text, heading_level). Supports simple Markdown headings."""
    blocks: list[tuple[str, str, int | None]] = []
    pending: list[str] = []

    def flush_pending() -> None:
        nonlocal pending
        if pending:
            blocks.append(("paragraph", "\n".join(pending).strip(), None))
            pending = []

    for raw in text.replace("\r\n", "\n").split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_pending()
            continue
        if stripped.startswith("#"):
            marker, _, rest = stripped.partition(" ")
            if marker and all(ch == "#" for ch in marker) and rest:
                flush_pending()
                level = min(len(marker), 6)
                blocks.append(("heading", rest.strip(), level))
                continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            flush_pending()
            blocks.append(("bullet", stripped[2:].strip(), None))
            continue
        pending.append(stripped)
    flush_pending()
    return blocks


def apply_text_to_run(run, text: str) -> None:
    if "\n" not in text:
        run.text = text
        return
    parts = text.split("\n")
    run.text = parts[0]
    for part in parts[1:]:
        run.add_break()
        extra = run._parent.add_run(part)
        copy_run_properties(run, extra)


def apply_block_to_paragraph(document, paragraph, block: tuple[str, str, int | None], template_paragraph=None) -> None:
    kind, text, level = block
    clear_paragraph_content(paragraph)

    if kind == "heading":
        style = heading_style(document, level or 1)
        if style:
            paragraph.style = style
    elif kind == "bullet":
        style = bullet_style(document)
        if style:
            paragraph.style = style
    else:
        if template_paragraph is not None:
            replace_paragraph_properties(template_paragraph, paragraph)
            if template_paragraph.style is not None:
                paragraph.style = template_paragraph.style
        else:
            style = paragraph_style(document)
            if style:
                paragraph.style = style

    run = paragraph.add_run()
    source_run = first_text_run(template_paragraph) if template_paragraph is not None and kind == "paragraph" else None
    copy_run_properties(source_run, run)
    apply_text_to_run(run, text)


def pick_template_near_index(document, anchor_idx: int) -> object | None:
    paragraphs = document.paragraphs
    for idx in range(anchor_idx + 1, len(paragraphs)):
        para = paragraphs[idx]
        if heading_level_for_paragraph(para) is None and not paragraph_is_protected(para):
            return para
    for idx in range(anchor_idx - 1, -1, -1):
        para = paragraphs[idx]
        if heading_level_for_paragraph(para) is None and not paragraph_is_protected(para):
            return para
    if 0 <= anchor_idx < len(paragraphs):
        return paragraphs[anchor_idx]
    return None


def add_blocks(document, blocks: list[tuple[str, str, int | None]]) -> int:
    template = None
    for para in reversed(document.paragraphs):
        if heading_level_for_paragraph(para) is None and not paragraph_is_protected(para):
            template = para
            break

    count = 0
    for kind, text, level in blocks:
        if not text:
            continue
        para = document.add_paragraph()
        apply_block_to_paragraph(document, para, (kind, text, level), template_paragraph=template)
        if kind == "paragraph":
            template = para
        count += 1
    return count


def insert_blocks_after(document, anchor_idx: int, blocks: list[tuple[str, str, int | None]]) -> int:
    anchor = document.paragraphs[anchor_idx]
    template = pick_template_near_index(document, anchor_idx)
    current = anchor
    count = 0
    for block in blocks:
        kind, text, _ = block
        if not text:
            continue
        current = insert_paragraph_after(current)
        apply_block_to_paragraph(document, current, block, template_paragraph=template)
        if kind == "paragraph":
            template = current
        count += 1
    return count


def find_heading(document, heading: str) -> int | None:
    target = normalized_text(heading)
    for idx, para in enumerate(document.paragraphs):
        text = normalized_text(para.text)
        if text == target:
            return idx
    return None


def section_body_bounds(document, heading_idx: int) -> tuple[int, int | None, list[int]]:
    paragraphs = document.paragraphs
    if not (0 <= heading_idx < len(paragraphs)):
        raise SystemExit("Heading index is out of range.")

    current_level = heading_level_for_paragraph(paragraphs[heading_idx])
    nested_heading_indexes: list[int] = []
    for idx in range(heading_idx + 1, len(paragraphs)):
        level = heading_level_for_paragraph(paragraphs[idx])
        if level is None:
            continue
        if current_level is not None and level > current_level:
            nested_heading_indexes.append(idx)
            continue
        return heading_idx + 1, idx - 1, nested_heading_indexes
    return heading_idx + 1, len(paragraphs) - 1, nested_heading_indexes


def protected_paragraph_indexes(document, start_idx: int, end_idx: int | None) -> list[int]:
    if end_idx is None or end_idx < start_idx:
        return []
    protected = []
    for idx in range(start_idx, end_idx + 1):
        if paragraph_is_protected(document.paragraphs[idx]):
            protected.append(idx)
    return protected


def editable_segments(document, start_idx: int, end_idx: int | None) -> list[tuple[int, int]]:
    if end_idx is None or end_idx < start_idx:
        return []
    segments: list[tuple[int, int]] = []
    current_start: int | None = None
    for idx in range(start_idx, end_idx + 1):
        if paragraph_is_protected(document.paragraphs[idx]):
            if current_start is not None:
                segments.append((current_start, idx - 1))
                current_start = None
            continue
        if current_start is None:
            current_start = idx
    if current_start is not None:
        segments.append((current_start, end_idx))
    return segments


def segment_slot_count(segment: tuple[int, int]) -> int:
    start_idx, end_idx = segment
    return max(0, end_idx - start_idx + 1)


def allocate_blocks_to_segments(segments: list[tuple[int, int]], block_count: int) -> list[int]:
    if not segments:
        return []
    remaining_blocks = block_count
    remaining_slots = sum(segment_slot_count(segment) for segment in segments)
    allocation: list[int] = []
    for idx, segment in enumerate(segments):
        slot_count = segment_slot_count(segment)
        future_segments = len(segments) - idx - 1
        if remaining_blocks <= 0:
            allocation.append(0)
            remaining_slots -= slot_count
            continue

        reserve_for_future = min(remaining_blocks - 1, future_segments) if future_segments > 0 else 0
        if remaining_slots <= 0:
            proposed = remaining_blocks
        else:
            proposed = round((remaining_blocks * slot_count) / remaining_slots)
        proposed = max(1, proposed)
        proposed = min(proposed, remaining_blocks - reserve_for_future)
        allocation.append(proposed)
        remaining_blocks -= proposed
        remaining_slots -= slot_count

    if remaining_blocks > 0:
        allocation[-1] += remaining_blocks
    return allocation


def replace_paragraph_range(
    document,
    start_idx: int,
    end_idx: int,
    blocks: list[tuple[str, str, int | None]],
) -> dict[str, object]:
    paragraphs = document.paragraphs
    if start_idx < 0 or end_idx < start_idx or end_idx >= len(paragraphs):
        raise SystemExit("Invalid paragraph range.")

    protected = protected_paragraph_indexes(document, start_idx, end_idx)
    if protected:
        raise SystemExit(
            f"Protected paragraphs in requested range: {protected}. "
            "Refusing to rewrite paragraphs that contain drawings or page breaks."
        )

    target = [paragraphs[idx] for idx in range(start_idx, end_idx + 1)]
    original_count = len(target)
    current_anchor = target[-1]

    changed = 0
    for idx, block in enumerate(blocks):
        kind, text, _ = block
        if not text:
            continue
        if idx < original_count:
            para = target[idx]
            template = target[idx]
        else:
            para = insert_paragraph_after(current_anchor)
            template = target[min(idx, original_count - 1)]
            current_anchor = para
        apply_block_to_paragraph(document, para, block, template_paragraph=template)
        if idx < original_count:
            current_anchor = para
        changed += 1

    for extra in reversed(target[len(blocks) :]):
        remove_paragraph(extra)

    return {
        "paragraphs_written": changed,
        "range_start": start_idx,
        "range_end": end_idx,
        "protected_paragraphs": protected,
    }


def append_blocks_after_paragraph(document, anchor_idx: int, blocks: list[tuple[str, str, int | None]]) -> int:
    if not blocks:
        return 0
    return insert_blocks_after(document, anchor_idx, blocks)


def replace_section_body_preserving_protected(
    document,
    heading_idx: int,
    start_idx: int,
    end_idx: int,
    blocks: list[tuple[str, str, int | None]],
) -> dict[str, object]:
    segments = editable_segments(document, start_idx, end_idx)
    protected = protected_paragraph_indexes(document, start_idx, end_idx)
    if not segments:
        last_idx = end_idx if end_idx is not None and end_idx >= start_idx else heading_idx
        written = append_blocks_after_paragraph(document, last_idx, blocks)
        return {
            "paragraphs_written": written,
            "range_start": start_idx,
            "range_end": end_idx,
            "protected_paragraphs": protected,
            "editable_segments": [],
            "preserved_media": True,
    }

    allocation = allocate_blocks_to_segments(segments, len(blocks))
    segment_blocks: list[list[tuple[str, str, int | None]]] = []
    cursor = 0
    for allocated in allocation:
        seg_blocks = blocks[cursor : cursor + allocated]
        cursor += allocated
        segment_blocks.append(seg_blocks)

    changed = 0
    for segment, seg_blocks in reversed(list(zip(segments, segment_blocks, strict=False))):
        seg_start, seg_end = segment
        result = replace_paragraph_range(document, seg_start, seg_end, seg_blocks)
        changed += int(result["paragraphs_written"])

    if cursor < len(blocks):
        last_segment_end = segments[-1][1]
        changed += append_blocks_after_paragraph(document, last_segment_end, blocks[cursor:])

    return {
        "paragraphs_written": changed,
        "range_start": start_idx,
        "range_end": end_idx,
        "protected_paragraphs": protected,
        "editable_segments": [[seg_start, seg_end] for seg_start, seg_end in segments],
        "preserved_media": bool(protected),
    }


def replace_section_body(document, heading: str, blocks: list[tuple[str, str, int | None]]) -> dict[str, object]:
    heading_idx = find_heading(document, heading)
    if heading_idx is None:
        raise SystemExit(f"Heading not found: {heading}")

    start_idx, end_idx, nested_headings = section_body_bounds(document, heading_idx)
    if nested_headings:
        raise SystemExit(
            f"Nested headings found in the target body range: {nested_headings}. "
            "Refusing broad replacement because it can break sub-section structure."
        )

    if end_idx is None or end_idx < start_idx:
        written = insert_blocks_after(document, heading_idx, blocks)
        return {
            "paragraphs_written": written,
            "range_start": start_idx,
            "range_end": start_idx - 1,
            "protected_paragraphs": [],
            "editable_segments": [],
            "preserved_media": False,
            "heading": heading,
        }

    result = replace_section_body_preserving_protected(document, heading_idx, start_idx, end_idx, blocks)
    result["heading"] = heading
    return result


def inspect_docx(docx_path: Path) -> dict:
    docx = ensure_docx()
    document = docx.Document(str(docx_path))
    outline = []
    table_summaries = []
    media_paragraphs = []
    page_break_paragraphs = []
    heading_ranges = []

    for idx, para in enumerate(document.paragraphs):
        style = para.style.name if para.style is not None else ""
        text = para.text.strip()
        flags = paragraph_guard_flags(para)
        if flags["has_drawing"]:
            media_paragraphs.append({"index": idx, "style": style, "text": text[:180]})
        if flags["has_page_break"]:
            page_break_paragraphs.append({"index": idx, "style": style, "text": text[:180]})
        level = heading_level_for_paragraph(para)
        if text and level is not None:
            outline.append({"index": idx, "level": level, "style": style, "text": text[:180]})

    for entry in outline:
        start_idx, end_idx, nested = section_body_bounds(document, entry["index"])
        protected = protected_paragraph_indexes(document, start_idx, end_idx)
        heading_ranges.append(
            {
                "heading_index": entry["index"],
                "heading_level": entry["level"],
                "heading_text": entry["text"],
                "body_start": start_idx,
                "body_end": end_idx,
                "nested_heading_indexes": nested,
                "protected_paragraphs": protected,
                "editable_segments": [list(segment) for segment in editable_segments(document, start_idx, end_idx)],
                "safe_body_replace": not nested,
                "preserves_media_during_replace": not nested,
            }
        )

    for idx, table in enumerate(document.tables, start=1):
        preview = ""
        if table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
            preview = " | ".join([cell for cell in cells if cell][:4])[:180]
        table_summaries.append(
            {
                "index": idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "first_row_preview": preview,
            }
        )

    header_texts = []
    footer_texts = []
    for idx, section in enumerate(document.sections, start=1):
        header = " ".join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])[:180]
        footer = " ".join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])[:180]
        header_texts.append({"section": idx, "text": header})
        footer_texts.append({"section": idx, "text": footer})

    return {
        "docx": str(docx_path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "inline_shape_count": len(document.inline_shapes),
        "style_count": len(document.styles),
        "style_samples": sorted(style_names(document))[:40],
        "outline": outline,
        "heading_ranges": heading_ranges,
        "media_paragraphs": media_paragraphs,
        "page_break_paragraphs": page_break_paragraphs,
        "table_summaries": table_summaries,
        "headers": header_texts,
        "footers": footer_texts,
        "guardrail": (
            "Use inspect first, then choose the smallest authorized edit surface that preserves "
            "surrounding layout and non-text anchors."
        ),
    }


def load_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def write_log(workspace: Path, payload: dict) -> Path:
    log_dir = state_root(workspace) / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    path = log_dir / f"{timestamp()}-docx-writer.json"
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def save_document(document, out_path: Path) -> None:
    document.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="command", required=True)

    p_inspect = sub.add_parser("inspect")
    p_inspect.add_argument("--docx", required=True)

    for name in ("append", "insert-after-heading", "replace-paragraph-range", "replace-section-body"):
        p = sub.add_parser(name)
        if name != "inspect":
            p.add_argument("--workspace", required=True)
            p.add_argument("--docx", required=True)
            p.add_argument("--text-file", required=True)
            p.add_argument("--save-as")
            p.add_argument("--label", default=name)
        if name in {"insert-after-heading", "replace-section-body"}:
            p.add_argument("--heading", required=True)
        if name == "replace-paragraph-range":
            p.add_argument("--start-paragraph", required=True, type=int)
            p.add_argument("--end-paragraph", required=True, type=int)

    args = parser.parse_args()
    docx = ensure_docx()

    if args.command == "inspect":
        print(json.dumps(inspect_docx(Path(args.docx).resolve()), ensure_ascii=False, indent=2))
        return 0

    workspace = Path(args.workspace).resolve()
    docx_path = resolve_workspace_path(workspace, args.docx)
    out_path = resolve_workspace_path(workspace, args.save_as) if args.save_as else docx_path
    text = load_text(resolve_workspace_path(workspace, args.text_file))
    blocks = parse_text_blocks(text)
    document = docx.Document(str(docx_path))
    backup = backup_docx(workspace, docx_path, args.label)

    if args.command == "append":
        result = {"paragraphs_written": add_blocks(document, blocks)}
    elif args.command == "insert-after-heading":
        idx = find_heading(document, args.heading)
        if idx is None:
            raise SystemExit(f"Heading not found: {args.heading}")
        result = {"paragraphs_written": insert_blocks_after(document, idx, blocks), "heading": args.heading}
    elif args.command == "replace-paragraph-range":
        result = replace_paragraph_range(document, args.start_paragraph, args.end_paragraph, blocks)
    elif args.command == "replace-section-body":
        result = replace_section_body(document, args.heading, blocks)
    else:  # pragma: no cover
        raise SystemExit(f"Unknown command: {args.command}")

    save_document(document, out_path)
    payload = {
        "command": args.command,
        "docx": str(docx_path),
        "output": str(out_path),
        "backup": str(backup),
        "guardrail": (
            "Review the edited range in Word. The correct primitive depends on the frozen plan "
            "scope, desired edit surface, and need to preserve nearby layout."
        ),
        **result,
    }
    payload["log"] = str(write_log(workspace, payload))
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
